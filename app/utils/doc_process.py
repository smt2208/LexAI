import io
import asyncio
from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from app.logging_config import logger
from app.config import settings
from fastapi import UploadFile, HTTPException
import tempfile
import os

class DocumentProcessor:
    """Document processor for PDF and DOCX text extraction."""
    
    @staticmethod
    async def extract_text(file: UploadFile) -> str:
        """Extract text from uploaded file with validation."""
        try:
            # Validate file type
            if not file.filename:
                raise HTTPException(
                    status_code=400,
                    detail="File must have a name"
                )
            
            filename = file.filename.lower()
            allowed_extensions = ['.pdf', '.docx']
            if not any(filename.endswith(ext) for ext in allowed_extensions):
                raise HTTPException(
                    status_code=400,
                    detail="Only PDF and DOCX files are supported"
                )
            
            # Validate content type
            content_type = file.content_type or ""
            allowed_content_types = [
                'application/pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/octet-stream'  # Some browsers send this for various files
            ]
            
            if content_type and content_type not in allowed_content_types:
                logger.warning(f"Unexpected content type: {content_type}")
            
            file_content = await file.read()
            
            # Validate file size
            file_size_mb = len(file_content) / (1024 * 1024)
            if file_size_mb > settings.MAX_FILE_SIZE_MB:
                raise HTTPException(
                    status_code=413,
                    detail=f"File size {file_size_mb:.1f}MB exceeds maximum allowed size of {settings.MAX_FILE_SIZE_MB}MB"
                )
            
            # Validate file is not empty
            if len(file_content) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="File is empty"
                )
            
            logger.info(f"Processing file: {file.filename} ({file_size_mb:.1f}MB)")
            
            if filename.endswith('.pdf'):
                return await DocumentProcessor._extract_from_pdf(file_content)
            elif filename.endswith('.docx'):
                return await DocumentProcessor._extract_from_docx(file_content)
                
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from document: {str(e)}"
            )
    
    @staticmethod
    async def _extract_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF using PyPDFLoader with proper resource cleanup."""
        temp_file_path = None
        try:
            # Create temporary file with context manager for better resource management
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            def load_pdf():
                loader = PyPDFLoader(temp_file_path)
                documents = loader.load()
                return "\n\n".join([doc.page_content for doc in documents])
            
            # Run PDF loading in executor with timeout
            text = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(None, load_pdf),
                timeout=30.0  # 30 second timeout
            )
            
            if not text.strip():
                raise ValueError("No readable text found in PDF")
                
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()
                
        except asyncio.TimeoutError:
            logger.error("PDF extraction timed out after 30 seconds")
            raise HTTPException(
                status_code=408,
                detail="PDF processing timed out. Please try with a smaller file."
            )
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from PDF: {str(e)}"
            )
        finally:
            # Ensure cleanup even if an exception occurs
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                    logger.debug(f"Cleaned up temporary file: {temp_file_path}")
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temporary file {temp_file_path}: {cleanup_error}")
    
    @staticmethod
    async def _extract_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX with timeout and better error handling."""
        try:
            def extract_docx():
                docx_file = io.BytesIO(file_content)
                doc = Document(docx_file)
                
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            text_parts.append(row_text)
                
                return "\n\n".join(text_parts)
            
            # Run DOCX extraction in executor with timeout
            text = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(None, extract_docx),
                timeout=20.0  # 20 second timeout
            )
            
            if not text.strip():
                raise ValueError("No readable text found in DOCX")
                
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except asyncio.TimeoutError:
            logger.error("DOCX extraction timed out after 20 seconds")
            raise HTTPException(
                status_code=408,
                detail="DOCX processing timed out. Please try with a smaller file."
            )
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from DOCX: {str(e)}"
            )
    
    @staticmethod
    def validate_file(file: UploadFile, max_size_mb: int = 10) -> None:
        """Validate uploaded file type and basic properties."""
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        filename = file.filename.lower()
        if not (filename.endswith('.pdf') or filename.endswith('.docx')):
            raise HTTPException(
                status_code=400, 
                detail="Only PDF and DOCX files are allowed"
            )
        
        logger.info(f"File validation passed: {file.filename}")